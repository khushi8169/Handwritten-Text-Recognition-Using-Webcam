import cv2
import numpy as np
import os
import tensorflow as tf
from tensorflow.keras.models import load_model
import matplotlib.pyplot as plt
from docx import Document

model = load_model("English.h5")

emnist_mapping = {0:'0',1:'1',2:'2',3:'3',4:'4',5:'5',6:'6',7:'7',8:'8',9:'9',10:'A',11:'B',12:'C',13:'D',14:'E',15:'F',
             16:'G',17:'H',18:'I',19:'J',20:'K',21:'L',22:'M',23:'N', 24:'O',25:'P',26:'Q',27:'R',28:'S',29:'T',30:'U',
             31:'V',32:'W',33:'X',34:'Y',35:'Z',36:'a',37:'b',38:'c',39:'d',40:'e',41:'f',42:'g',43:'h',44:'i',45:'j',
             46:'k',47:'l',48:'m',49:'n', 50:'o',51:'p',52:'q',53:'r',54:'s',55:'t',56:'u',57:'v',58:'w',59:'x',60:'y',
             61:'z'}

def preprocess_image(image_path, debug=False):
    image = cv2.imread(image_path)
    if image is None:
        raise ValueError(f"Could not read image at {image_path}")
    
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    
    expanded_img = cv2.copyMakeBorder(gray, 20, 20, 20, 20, cv2.BORDER_CONSTANT, value=255)
    
    blurred = cv2.GaussianBlur(expanded_img, (5, 5), 0)
    
    thresh = cv2.adaptiveThreshold(blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                  cv2.THRESH_BINARY_INV, 11, 2)
    
    kernel = np.ones((3, 3), np.uint8)
    thresh = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
    
    if debug:
        plt.figure(figsize=(15, 5))
        plt.subplot(131), plt.imshow(gray, cmap='gray'), plt.title('Original')
        plt.subplot(132), plt.imshow(blurred, cmap='gray'), plt.title('Blurred')
        plt.subplot(133), plt.imshow(thresh, cmap='gray'), plt.title('Thresholded')
        plt.tight_layout()
        plt.savefig("preprocessing_debug.png")
        plt.close()
    
    return thresh

def segment_image(image, output_dir="segments", debug=False):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    contours, _ = cv2.findContours(image.copy(), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    
    valid_contours = []
    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        if w > 8 and h > 8 and w*h > 100: 
            valid_contours.append(cnt)
    
    valid_contours = sorted(valid_contours, key=lambda c: cv2.boundingRect(c)[0])
    
    segments = []
    debug_img = cv2.cvtColor(image.copy(), cv2.COLOR_GRAY2BGR) if debug else None
    
    for i, cnt in enumerate(valid_contours):
        x, y, w, h = cv2.boundingRect(cnt)
        
        if debug:
            cv2.rectangle(debug_img, (x, y), (x+w, y+h), (0, 255, 0), 2)
        
        mask = np.zeros_like(image)
        cv2.drawContours(mask, [cnt], -1, 255, thickness=cv2.FILLED)  # Fill the contour

        segment = mask[y:y+h, x:x+w]
        
        target_size = 28
        
        padding = int(max(w, h) * 0.2)
        padded_segment = cv2.copyMakeBorder(segment, padding, padding, padding, padding, 
                                           cv2.BORDER_CONSTANT, value=0)
        
        resized_segment = cv2.resize(padded_segment, (target_size, target_size))
        
        segment_path = os.path.join(output_dir, f"segment_{i}.png")
        cv2.imwrite(segment_path, resized_segment)
        
        
        segments.append(resized_segment)
    
    if debug and len(valid_contours) > 0:
        cv2.imwrite("segmentation_debug.png", debug_img)
    
    return segments

def predict_character(segment):

    img_gray = cv2.resize(segment, (28, 28))
    img_reshaped = img_gray.reshape(1, 28, 28, 1).astype('float32')
    pred = model.predict(img_reshaped)
    ans = emnist_mapping[np.argmax(pred)]
    
    return ans, np.max(pred) * 100  

def post_process_text(predicted_text):
    corrected_text = ""
    
    for i, char in enumerate(predicted_text):
        if char == "0":
            if (i > 0 and predicted_text[i - 1].isalpha()) or (i < len(predicted_text) - 1 and predicted_text[i + 1].isalpha()):
                corrected_text += "o" 
            else:
                corrected_text += "0"
        
        else:
            corrected_text += char

    return corrected_text

def visualize_predictions(segments, predictions):
    plt.figure(figsize=(15, 3))
    for i, (segment, (char, conf)) in enumerate(zip(segments, predictions)):
        plt.subplot(1, len(segments), i + 1)
        plt.imshow(segment, cmap='gray')
        plt.title(f"{char}\n{conf:.1f}%", fontsize=10)
        plt.axis('off')
    plt.tight_layout()
    plt.show()

def save_recognized_text(recognized_text, avg_confidence, output_path="recognized_output.docx"):
    doc = Document()

    doc.add_heading("Handwritten Sentence Recognition Results", level=1)

    doc.add_paragraph("Here's your recognized text without OCR, carefully analyzed with deep learning!\n")

    doc.add_heading("Recognized Text:", level=2)
    doc.add_paragraph(recognized_text, style="Normal") 

    doc.add_heading("Average Confidence Score:", level=2)
    doc.add_paragraph(f"{avg_confidence:.2f}%", style="Normal")

    doc.save(output_path)
    print(f"Recognition results saved to {output_path}")

if __name__ == "__main__":
    # image_path = "C:\\Users\\hp\\OneDrive\\Desktop\\miner\\captured_image.png"
    image_path = input("Enter the path of the image: ")
    debug_mode = 'y'
    
    processed_image = preprocess_image(image_path, debug=debug_mode)
    segments = segment_image(processed_image, debug=debug_mode)

    recognized_text = ""
    confidences = []
    predictions = []

    for segment in segments:
        char, conf = predict_character(segment)
        recognized_text += char
        confidences.append(conf)
        predictions.append((char, conf))

    avg_confidence = sum(confidences) / len(confidences) if confidences else 0
    corrected_text = post_process_text(recognized_text)

    visualize_predictions(segments, predictions)
    save_recognized_text(corrected_text, avg_confidence)
    
    print(f"Recognized Text: {corrected_text}")
    print(f"Average Confidence: {avg_confidence:.2f}%")
    
    print(f"Text saved to recognized_text.txt")